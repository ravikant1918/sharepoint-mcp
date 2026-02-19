"""File content parsers for PDF, Excel, and Word documents."""
from __future__ import annotations

import io
import logging
from typing import Literal, Tuple

logger = logging.getLogger(__name__)

# Mapping of logical type → file extensions
_FILE_TYPE_MAP: dict[str, list[str]] = {
    "text": [".txt", ".csv", ".json", ".xml", ".html", ".htm", ".md", ".js", ".ts", ".css", ".py", ".yaml", ".yml"],
    "pdf": [".pdf"],
    "excel": [".xlsx", ".xls"],
    "word": [".docx", ".doc"],
}

FileType = Literal["text", "pdf", "excel", "word", "binary"]


def detect_file_type(file_name: str) -> FileType:
    """Return the logical file type for *file_name* based on its extension."""
    lower = file_name.lower()
    for ftype, extensions in _FILE_TYPE_MAP.items():
        if any(lower.endswith(ext) for ext in extensions):
            return ftype  # type: ignore[return-value]
    return "binary"


def parse_pdf(content_bytes: bytes) -> Tuple[str, int]:
    """Extract plain text from a PDF byte string.

    Returns:
        (text, page_count)

    Raises:
        Exception: propagates any fitz / PyMuPDF error.
    """
    import fitz  # PyMuPDF — optional heavy dep, imported lazily

    doc = fitz.open(stream=content_bytes, filetype="pdf")
    text = "".join(doc[i].get_text() + "\n" for i in range(len(doc)))
    page_count = len(doc)
    doc.close()
    return text.strip(), page_count


def parse_excel(content_bytes: bytes) -> Tuple[str, int]:
    """Extract sheet data from an Excel file.

    Returns:
        (text, sheet_count)
    """
    import pandas as pd  # optional heavy dep, imported lazily

    sheets: dict = pd.read_excel(io.BytesIO(content_bytes), sheet_name=None)
    parts: list[str] = []
    for sheet_name, df in sheets.items():
        parts.append(f"=== {sheet_name} ===")
        parts.extend(
            df.head(50).fillna("").astype(str).apply(" | ".join, axis=1).tolist()
        )
    return "\n".join(parts), len(sheets)


def parse_word(content_bytes: bytes) -> Tuple[str, int]:
    """Extract text from a Word document.

    Returns:
        (text, paragraph_count)
    """
    from docx import Document  # python-docx, imported lazily

    doc = Document(io.BytesIO(content_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts), len(doc.paragraphs)
